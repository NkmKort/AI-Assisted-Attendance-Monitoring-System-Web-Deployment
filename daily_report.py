import os
import io
import smtplib
import pandas as pd
from datetime import datetime
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from docx import Document
from supabase import create_client, Client

def get_supabase() -> Client:
    """Helper to initialize Supabase client with URL normalization."""
    url = os.environ.get("SUPABASE_URL")
    key = os.environ.get("SUPABASE_KEY")
    if not url or not key:
        raise ValueError("Missing SUPABASE_URL or SUPABASE_KEY environment variables.")
    
    # Normalize URL (remove trailing /v1 or /rest/v1)
    if "/rest/v1" in url:
        url = url.split("/rest/v1")[0]
    elif "/v1" in url:
        url = url.split("/v1")[0]
    url = url.rstrip("/")
    
    return create_client(url, key)

def _normalize_bools(df: pd.DataFrame) -> pd.DataFrame:
    for col in ["lanyard_compliant", "dresscode_compliant", "synced"]:
        if col in df.columns:
            df[col] = df[col].apply(
                lambda v: str(v).strip().lower() in ("true", "1", "yes", "t")
                if not isinstance(v, bool) else v
            )
    return df

def load_and_prepare(csv_file=None) -> pd.DataFrame:
    """Loads attendance data from Supabase and prepares it for the report generator."""
    try:
        sb = get_supabase()
        resp = (
            sb.table("attendance")
            .select("student_name, timestamp, synced, lanyard_compliant, dresscode_compliant")
            .order("timestamp", desc=True)
            .limit(2000)
            .execute()
        )
        if not resp.data:
            return pd.DataFrame(columns=[
                "Student_Name", "Timestamp", "Cloud_Synced",
                "Lanyard_Compliant", "DressCode_Compliant", "_date"
            ])
            
        df = pd.DataFrame(resp.data)
        df = _normalize_bools(df)
        df["timestamp"] = pd.to_datetime(df["timestamp"], utc=True, errors="coerce")
        df = df.sort_values("timestamp", ascending=False).reset_index(drop=True)
        df = df.rename(columns={
            "student_name":       "Student_Name",
            "timestamp":          "Timestamp",
            "synced":             "Cloud_Synced",
            "lanyard_compliant":  "Lanyard_Compliant",
            "dresscode_compliant":"DressCode_Compliant",
        })
        
        # Add _date column (date objects) for date filter matching
        df["_date"] = df["Timestamp"].dt.date
        return df
    except Exception as e:
        print(f"Could not load attendance from Supabase in daily_report: {e}")
        return pd.DataFrame(columns=[
            "Student_Name", "Timestamp", "Cloud_Synced",
            "Lanyard_Compliant", "DressCode_Compliant", "_date"
        ])

def filter_by_date(df: pd.DataFrame, selected_date) -> pd.DataFrame:
    """Filters data for the selected date."""
    if df.empty:
        return df
    return df[df["_date"] == selected_date]

def build_stats(df: pd.DataFrame) -> dict:
    """Builds count statistics of compliance and violations."""
    if df.empty:
        return {
            "total": 0,
            "compliant": 0,
            "lanyard_violations": 0,
            "dress_violations": 0
        }
    return {
        "total": len(df),
        "compliant": int((df["Lanyard_Compliant"] & df["DressCode_Compliant"]).sum()),
        "lanyard_violations": int((~df["Lanyard_Compliant"]).sum()),
        "dress_violations": int((~df["DressCode_Compliant"]).sum())
    }

def generate_report(llm, vector_store, selected_date, recipient_email, daily_df):
    """Retrieves context policies from RAG vector store and compiles a formal report & email draft."""
    query = "ID compliance, lanyard rules, dress code guidelines and policy actions"
    try:
        docs = vector_store.similarity_search(query, k=3)
        context = "\n\n".join(doc.page_content for doc in docs)
    except Exception as e:
        context = "No specific student handbook policies retrieved."
        
    stats = build_stats(daily_df)
    
    # Format the violators details
    violators_df = daily_df[~(daily_df["Lanyard_Compliant"] & daily_df["DressCode_Compliant"])]
    violators_list = []
    for _, row in violators_df.iterrows():
        issues = []
        if not row["Lanyard_Compliant"]:
            issues.append("Missing ID/Lanyard")
        if not row["DressCode_Compliant"]:
            issues.append("Dress Code Violation")
        violators_list.append(f"- {row['Student_Name']}: {', '.join(issues)}")
    violators_text = "\n".join(violators_list) if violators_list else "No violations recorded today."
    
    prompt = f"""You are the official AI Assistant for the Mapúa University Student Handbook.
    Your task is to write a formal Daily Compliance Report and an Email Draft based on the attendance logs for {selected_date}.
    
    DAILY STATISTICS:
    - Total Students Logged: {stats['total']}
    - Fully Compliant Students: {stats['compliant']}
    - Missing ID/Lanyard Violations: {stats['lanyard_violations']}
    - Dress Code Violations: {stats['dress_violations']}
    
    LIST OF STUDENTS WITH VIOLATIONS:
    {violators_text}
    
    RELEVANT UNIVERSITY HANDBOOK POLICIES:
    {context}
    
    Please write:
    1. A formal, detailed administrative report addressed to the Discipline Office, summarizing the statistics, listing the violators, and referencing relevant handbook policies and penalties.
    2. An email draft to the Discipline Office ({recipient_email}) presenting this report.
    
    IMPORTANT: You MUST separate the two drafts with exactly this delimiter: '---EMAIL_DRAFT---'
    """
    
    response = llm.invoke(prompt).content
    
    if "---EMAIL_DRAFT---" in response:
        word_report, email_draft = response.split("---EMAIL_DRAFT---", 1)
    else:
        word_report = response
        email_draft = f"Subject: Daily Compliance Report - {selected_date}\n\nDear Discipline Office,\n\nPlease find attached the Daily Compliance Report for {selected_date}.\n\nBest regards,\nFateAutomata Kiosk"
        
    return word_report.strip(), email_draft.strip()

def export_to_docx(report_text: str, report_date) -> bytes:
    """Compiles the report text into a Word docx structure."""
    doc = Document()
    doc.add_heading("FateAutomata Daily Compliance Report", level=0)
    doc.add_paragraph(f"Report Date: {report_date}")
    doc.add_paragraph("=" * 50)
    
    for p_text in report_text.split("\n\n"):
        if p_text.strip():
            doc.add_paragraph(p_text.strip())
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def send_email(recipient, subject, body, docx_bytes, docx_filename):
    """Sends the report and attachment via SMTP (requires app password)."""
    sender = os.environ.get("SMTP_SENDER")
    password = os.environ.get("SMTP_PASSWORD")
    
    if not sender or not password:
        raise ValueError("SMTP credentials missing. Please define SMTP_SENDER and SMTP_PASSWORD in your .env file.")
        
    # Set up message
    msg = MIMEMultipart()
    msg['From'] = sender
    msg['To'] = recipient
    msg['Subject'] = subject
    
    msg.attach(MIMEText(body, 'plain'))
    
    # Attach file
    part = MIMEBase('application', 'octet-stream')
    part.set_payload(docx_bytes)
    encoders.encode_base64(part)
    part.add_header('Content-Disposition', f"attachment; filename= {docx_filename}")
    msg.attach(part)
    
    # Connect and send
    server = smtplib.SMTP('smtp.gmail.com', 587)
    server.starttls()
    server.login(sender, password)
    server.sendmail(sender, recipient, msg.as_string())
    server.quit()
